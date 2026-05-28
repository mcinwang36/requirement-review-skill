#!/usr/bin/env python3
"""
AI 需求共创产品经理 Agent — 评审报告生成器
动态聚焦分析结构，兼容旧版评分结构
"""

import sys
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Color palette ──
COLOR_PRIMARY   = "1A3C6E"
COLOR_SECONDARY = "4A6FA5"
COLOR_ACCENT    = "2E75B6"
COLOR_PASS      = "2D7D46"
COLOR_WARN      = "B8860B"
COLOR_FAIL      = "8B3A3A"
COLOR_BODY      = "333333"
COLOR_MUTED     = "888888"
COLOR_BORDER    = "D0D0D0"
COLOR_HEADER_BG = "1A3C6E"
COLOR_ALT_ROW   = "F5F7FA"


def hex_rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_bg(cell, hex_color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=Pt(9.5), align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align:
        para.alignment = align
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(str(text))
    run.font.size = size
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = hex_rgb(COLOR_PRIMARY if level <= 1 else COLOR_SECONDARY)
    return h


def add_para(doc, text, bold=False, indent=None, color=None, size=Pt(10.5)):
    para = doc.add_paragraph()
    run = para.add_run(str(text))
    run.font.size = size
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = hex_rgb(color or COLOR_BODY)
    if bold:
        run.bold = True
    if indent:
        para.paragraph_format.left_indent = Cm(indent)
    para.paragraph_format.space_after = Pt(4)
    return para


def add_divider(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{COLOR_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def build_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{COLOR_BORDER}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(cell, COLOR_HEADER_BG)
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if isinstance(cell_data, tuple):
                text, bold, color = cell_data
            else:
                text, bold, color = str(cell_data), False, hex_rgb(COLOR_BODY)
            set_cell_text(cell, text, bold=bold, color=color)
            if r_idx % 2 == 1:
                set_cell_bg(cell, COLOR_ALT_ROW)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def render_core_contradiction(doc, data):
    cc = data.get("core_contradiction")
    if not cc:
        return
    add_heading(doc, "核心矛盾", 2)
    ctype = cc.get("type", "")
    desc = cc.get("description", "")
    why = cc.get("why_it_matters", "")
    if ctype:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(f"矛盾类型：{ctype}")
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.color.rgb = hex_rgb(COLOR_FAIL)
    if desc:
        add_para(doc, desc, indent=0.5, size=Pt(10.5))
    if why:
        add_para(doc, f"如不解决：{why}", indent=0.5, color=COLOR_WARN, size=Pt(10))


def render_ai_necessity(doc, data):
    necessity = data.get("ai_necessity", "")
    reason = data.get("ai_necessity_reason", "")
    if not necessity:
        return
    add_heading(doc, "AI 必要性", 2)
    color = COLOR_PASS if necessity == "必要" else (COLOR_WARN if "部分" in necessity else COLOR_FAIL)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f"判断：{necessity}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.color.rgb = hex_rgb(color)
    if reason:
        add_para(doc, reason, indent=0.5, size=Pt(10))


def render_next_steps(doc, data):
    ns = data.get("next_steps")
    if not ns:
        return

    must_do = ns.get("must_do_before_start", [])
    if must_do:
        add_para(doc, "立项前必须解决：", bold=True, size=Pt(10.5))
        for item in must_do:
            add_para(doc, f"  → {item}", indent=0.5, color=COLOR_FAIL, size=Pt(10))

    mvp = ns.get("mvp_scope", "")
    if mvp:
        add_para(doc, "MVP 范围：", bold=True, size=Pt(10.5))
        add_para(doc, mvp, indent=0.5, size=Pt(10))

    phases = ns.get("phase_plan", [])
    if phases:
        add_para(doc, "分阶段计划：", bold=True, size=Pt(10.5))
        headers = ["阶段", "重点", "成功标准"]
        rows = []
        for ph in phases:
            rows.append([
                ph.get("phase", "-"),
                ph.get("focus", "-"),
                ph.get("success_criteria", "-"),
            ])
        build_table(doc, headers, rows, col_widths=[2.5, 6, 5])

    defer = ns.get("defer_to_later", [])
    if defer:
        add_para(doc, "当前阶段暂缓：", bold=True, size=Pt(10.5))
        for item in defer:
            add_para(doc, f"  ✗ {item}", indent=0.5, color=COLOR_MUTED, size=Pt(10))


def render_focused_analysis(doc, data):
    """渲染动态聚焦分析（新版），兼容旧版 dimension_reviews"""
    analyses = data.get("focused_analysis", [])
    if not analyses:
        # 兼容旧版
        dim_reviews = data.get("dimension_reviews", {})
        if dim_reviews:
            _render_dimension_reviews_legacy(doc, dim_reviews)
        return

    for item in analyses:
        if not isinstance(item, dict):
            continue
        focus = item.get("focus", "")
        current = item.get("current_state", "")
        problem = item.get("core_problem", "")
        risk = item.get("risk", "")
        rec = item.get("recommendation", "")

        add_heading(doc, focus, 2)
        if current:
            add_para(doc, f"现状：{current}", indent=0.3, size=Pt(10))
        if problem:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            r = p.add_run(f"核心问题：{problem}")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.name = "微软雅黑"
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.color.rgb = hex_rgb(COLOR_FAIL)
        if risk:
            add_para(doc, f"风险：{risk}", indent=0.3, color=COLOR_WARN, size=Pt(10))
        if rec:
            add_para(doc, f"建议：{rec}", indent=0.3, color=COLOR_SECONDARY, size=Pt(10))


def render_issue_block(doc, issues, level_label, level_color):
    if not issues:
        return
    add_heading(doc, level_label, 2)
    for i, issue in enumerate(issues, 1):
        if isinstance(issue, dict):
            dim = issue.get("dimension", "")
            iss = issue.get("issue", issue.get("risk", issue.get("suggestion", "")))
            impact = issue.get("impact", "")
            action = issue.get("required_action", issue.get("mitigation", issue.get("suggestion", "")))
            likelihood = issue.get("likelihood", "")

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            label = f"[{dim}] " if dim else ""
            likelihood_str = f"（{likelihood}）" if likelihood else ""
            r = p.add_run(f"{i}. {label}{iss}{likelihood_str}")
            r.bold = True
            r.font.size = Pt(10.5)
            r.font.name = "微软雅黑"
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            r.font.color.rgb = hex_rgb(level_color)

            if impact:
                add_para(doc, f"影响：{impact}", indent=0.8, color=COLOR_BODY, size=Pt(10))
            if action:
                add_para(doc, f"建议：{action}", indent=0.8, color=COLOR_SECONDARY, size=Pt(10))
        else:
            add_para(doc, f"{i}. {issue}", indent=0.3)


def _render_dimension_reviews_legacy(doc, dim_reviews):
    """旧版五维评审兼容渲染"""
    dim_map = {
        "business_value":    ("业务价值", ""),
        "ai_necessity":      ("AI 必要性", ""),
        "capability_layers": ("能力分层", ""),
        "data_governance":   ("数据治理", ""),
        "feasibility":       ("可落地性", ""),
    }
    field_labels = {
        "strengths": ("优势", COLOR_PASS),
        "gaps": ("不足", COLOR_FAIL),
        "available": ("已有", COLOR_PASS),
        "missing": ("缺失", COLOR_FAIL),
        "critical_gap": ("关键缺口", COLOR_FAIL),
        "appropriate_ai_use": ("适合AI", COLOR_PASS),
        "over_engineering": ("过度设计", COLOR_WARN),
        "rule_layer": ("Rule层", COLOR_BODY),
        "workflow_layer": ("Workflow层", COLOR_BODY),
        "llm_layer": ("LLM层", COLOR_BODY),
        "feasible": ("可行", COLOR_PASS),
    }
    for key, (title, _) in dim_map.items():
        d = dim_reviews.get(key)
        if not d:
            continue
        add_heading(doc, title, 2)
        if isinstance(d, dict):
            conclusion = d.get("conclusion", "")
            if conclusion:
                add_para(doc, f"结论：{conclusion}", bold=True, indent=0.3, size=Pt(10.5))
            if key == "feasibility" and "gaps" in d:
                d["gaps_feasibility"] = d.pop("gaps")
            for field, (label, color) in field_labels.items():
                val = d.get(field)
                if not val:
                    continue
                items = val if isinstance(val, list) else [val]
                for item in items:
                    add_para(doc, f"  ▸ [{label}] {item}", color=color, size=Pt(10), indent=0.5)
        elif isinstance(d, str):
            add_para(doc, d, indent=0.3)


def generate_report_new(data: dict, output_path: str):
    """新版动态聚焦评审报告"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading('AI 需求评审意见', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = hex_rgb(COLOR_PRIMARY)
        run.font.size = Pt(22)

    meta = doc.add_paragraph()
    r1 = meta.add_run(f"评审日期：{datetime.now().strftime('%Y年%m月%d日')}")
    r1.font.size = Pt(9)
    r1.font.color.rgb = hex_rgb(COLOR_MUTED)
    r2 = meta.add_run("     |     密级：内部")
    r2.font.size = Pt(9)
    r2.font.color.rgb = hex_rgb(COLOR_MUTED)
    meta.paragraph_format.space_after = Pt(4)
    add_divider(doc)

    # 一、评审结论
    add_heading(doc, "一、评审结论", 1)

    verdict = data.get("verdict", "")
    summary = data.get("executive_summary", "")
    req_types = data.get("requirement_type", [])

    verdict_color = COLOR_PASS if "推进" in verdict else (COLOR_WARN if "条件" in verdict else COLOR_FAIL)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"总体判断：{verdict}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    r.font.color.rgb = hex_rgb(verdict_color)

    if req_types:
        types_str = " / ".join(req_types) if isinstance(req_types, list) else str(req_types)
        add_para(doc, f"需求类型：{types_str}", color=COLOR_MUTED, size=Pt(9.5))

    if summary:
        add_para(doc, summary, size=Pt(11))

    render_core_contradiction(doc, data)
    render_ai_necessity(doc, data)

    # 问题概览
    blocking = data.get("blocking_issues", [])
    main_risks = data.get("main_risks", data.get("high_risks", []))
    medium = data.get("medium_risks", [])
    opts = data.get("optimizations", [])

    if any([blocking, main_risks, medium, opts]):
        add_heading(doc, "问题概览", 2)
        headers = ["类别", "数量"]
        rows = []
        if blocking:
            rows.append([("阻塞问题", True, hex_rgb(COLOR_FAIL)), str(len(blocking))])
        if main_risks:
            rows.append([("主要风险", True, hex_rgb(COLOR_WARN)), str(len(main_risks))])
        if medium:
            rows.append([("次要风险", False, hex_rgb(COLOR_BODY)), str(len(medium))])
        if opts:
            rows.append([("优化建议", False, hex_rgb(COLOR_MUTED)), str(len(opts))])
        build_table(doc, headers, rows, col_widths=[4, 2])

    add_divider(doc)

    # 二、推进路径
    add_heading(doc, "二、推进路径", 1)
    render_next_steps(doc, data)
    add_divider(doc)

    # 三、聚焦分析
    add_heading(doc, "三、聚焦分析", 1)
    render_focused_analysis(doc, data)
    add_divider(doc)

    # 四、问题清单
    if any([blocking, main_risks, medium, opts]):
        add_heading(doc, "四、问题清单", 1)
        render_issue_block(doc, blocking, "🔴 阻塞问题", COLOR_FAIL)
        render_issue_block(doc, main_risks, "🟠 主要风险", COLOR_WARN)
        render_issue_block(doc, medium, "🟡 次要风险", COLOR_ACCENT)
        render_issue_block(doc, opts, "💡 优化建议", COLOR_MUTED)
        add_divider(doc)

    # 页脚
    doc.add_paragraph()
    add_divider(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 本评审意见由 AI 需求共创产品经理 Agent 自动生成 —")
    fr.font.size = Pt(8)
    fr.font.color.rgb = hex_rgb(COLOR_MUTED)
    fr.italic = True

    doc.save(output_path)
    print(f"评审报告已生成: {output_path}")


def generate_report_legacy(data: dict, output_path: str):
    """旧版评分结构兼容报告"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading('业务需求文档评审意见', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = hex_rgb(COLOR_PRIMARY)
        run.font.size = Pt(22)

    meta = doc.add_paragraph()
    r1 = meta.add_run(f"评审日期：{datetime.now().strftime('%Y年%m月%d日')}")
    r1.font.size = Pt(9)
    r1.font.color.rgb = hex_rgb(COLOR_MUTED)
    meta.paragraph_format.space_after = Pt(4)
    add_divider(doc)

    add_heading(doc, "一、评审概要", 1)
    total_score = data.get("total_score", 0)
    max_score = data.get("max_score", 100)
    passed = total_score >= 60
    result_para = doc.add_paragraph()
    r = result_para.add_run("评审结论：建议通过" if passed else "评审结论：建议驳回修订")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = hex_rgb(COLOR_PASS if passed else COLOR_FAIL)
    add_para(doc, f"综合评分：{total_score} / {max_score}（通过线 ≥ 60 分）")

    dimension_scores = data.get("dimensions", {})
    if dimension_scores:
        add_heading(doc, "各维度评估概览", 2)
        headers = ["评估维度", "满分", "得分", "达成率", "评价"]
        rows = []
        for dim_name, dim_data in dimension_scores.items():
            max_s = dim_data.get("max", 0)
            score = dim_data.get("score", 0)
            rate = f"{score/max_s*100:.0f}%" if max_s > 0 else "-"
            rating = dim_data.get("rating", "-")
            rows.append([dim_name, str(max_s), str(score), rate, rating])
        build_table(doc, headers, rows, col_widths=[3.5, 1.5, 1.5, 1.5, 2.5])

    summary = data.get("executive_summary", "")
    if summary:
        add_heading(doc, "摘要", 2)
        add_para(doc, summary)

    add_divider(doc)
    add_heading(doc, "二、逐项评审意见", 1)

    analyses = data.get("analyses", {})
    dim_order = data.get("dim_order", list(analyses.keys()))
    for dim_name in dim_order:
        if dim_name not in analyses:
            continue
        analysis = analyses[dim_name]
        dim_score = dimension_scores.get(dim_name, {})
        max_s = dim_score.get("max", 0)
        score = dim_score.get("score", 0)
        rating = dim_score.get("rating", "-")
        add_heading(doc, dim_name, 2)
        p = doc.add_paragraph()
        r = p.add_run(f"得分：{score}/{max_s}    评价：{rating}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = hex_rgb(COLOR_SECONDARY)
        if isinstance(analysis, dict):
            for obs in analysis.get("observations", []):
                add_para(doc, f"• {obs}", indent=0.5)
            for sug in analysis.get("suggestions", []):
                add_para(doc, f"→ {sug}", indent=0.5)
        elif isinstance(analysis, str):
            for line in analysis.strip().split("\n"):
                if line.strip():
                    add_para(doc, line.strip())

    add_divider(doc)
    add_heading(doc, "三、关键发现与改进建议", 1)

    issues = data.get("issues", [])
    if issues:
        add_heading(doc, "需关注事项", 2)
        headers = ["编号", "位置", "问题描述", "影响", "关注级别"]
        rows = []
        for i, issue in enumerate(issues, 1):
            rows.append([str(i), issue.get("location", "-"), issue.get("description", "-"),
                         issue.get("impact", "-"), issue.get("severity", "建议关注")])
        build_table(doc, headers, rows, col_widths=[1, 2.5, 5, 3, 2])

    suggestions_list = data.get("suggestions", [])
    if suggestions_list:
        add_heading(doc, "改进建议", 2)
        for i, sug in enumerate(suggestions_list, 1):
            if isinstance(sug, dict):
                p = doc.add_paragraph()
                r = p.add_run(f"{i}. {sug.get('title', '')}")
                r.bold = True
                r.font.size = Pt(10.5)
                r.font.name = "微软雅黑"
                r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                if sug.get("detail"):
                    add_para(doc, sug["detail"], indent=0.5)
            else:
                add_para(doc, f"{i}. {sug}")

    doc.add_paragraph()
    add_divider(doc)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 本评审意见由需求评审系统自动生成 —")
    fr.font.size = Pt(8)
    fr.font.color.rgb = hex_rgb(COLOR_MUTED)
    fr.italic = True

    doc.save(output_path)
    print(f"评审报告已生成: {output_path}")


def generate_report(data: dict, output_path: str):
    """自动判断新旧格式"""
    if "verdict" in data or "blocking_issues" in data or "focused_analysis" in data or "core_contradiction" in data:
        generate_report_new(data, output_path)
    else:
        generate_report_legacy(data, output_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python generate_report.py --file <JSON文件路径> [输出路径]")
        sys.exit(1)

    output_path = "评审意见.docx"

    if sys.argv[1] == "--file":
        json_path = sys.argv[2]
        output_path = sys.argv[3] if len(sys.argv) > 3 else output_path
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = json.loads(sys.argv[1])
        output_path = sys.argv[2] if len(sys.argv) > 2 else output_path

    generate_report(data, output_path)
