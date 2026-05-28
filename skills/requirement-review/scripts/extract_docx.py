#!/usr/bin/env python3
"""
Extract all text content and embedded images from a Word (.docx) document.
Robust extraction covering paragraphs, tables, headers, footers, and inline images.
Outputs: content.md (structured text with image markers) and images/ directory.
"""

import sys
import os
import json
import zipfile
import io
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
}


def extract_all_images_from_docx(docx_path: str, img_dir: Path) -> dict:
    """Extract ALL images from a docx file by traversing the zip archive directly.
    Returns a dict mapping relationship rId -> filename."""
    image_map = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if name.startswith('word/media/'):
                # Extract the image
                fname = Path(name).name
                # Try to find the relationship ID by looking at rels files
                z.extract(name, img_dir.parent)
                # Move to images directory
                src = img_dir.parent / name
                dest = img_dir / fname
                if src.exists() and not dest.exists():
                    import shutil
                    shutil.move(str(src), str(dest))

    # Also extract via python-docx relationships (for rId mapping)
    doc = Document(docx_path)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            ext = rel.target_ref.split(".")[-1].split("?")[0]
            if ext.lower() not in ('jpg', 'jpeg', 'png', 'gif', 'bmp', 'webp', 'tiff', 'tif', 'emf', 'wmf', 'svg'):
                ext = 'png'
            fname = f"image_{rel.rId}.{ext}"
            fpath = img_dir / fname
            if not fpath.exists():
                try:
                    with open(fpath, "wb") as f:
                        f.write(rel.target_part.blob)
                except Exception:
                    pass
            image_map[rel.rId] = fname

    return image_map


def find_image_rId_in_element(elem) -> str | None:
    """Find image relationship ID in an XML element (paragraph, table cell, etc.)."""
    ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

    # Check for blip elements (images in drawings)
    blips = elem.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    for blip in blips:
        rId = blip.get(f'{ns}embed')
        if rId:
            return rId

    # Check for pict elements (VML images - older format)
    imagedatas = elem.findall('.//' + '{urn:schemas-microsoft-com:vml}imagedata')
    for imagedata in imagedatas:
        rId = imagedata.get(f'{ns}id')
        if rId:
            return rId

    # Check for drawing elements
    drawings = elem.findall('.//' + '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
    for drawing in drawings:
        blips = drawing.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        for blip in blips:
            rId = blip.get(f'{ns}embed')
            if rId:
                return rId

    return None


def extract_paragraphs(doc: Document, image_map: dict) -> list[str]:
    """Extract all paragraphs, marking image positions."""
    lines = []
    para_index = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        has_image = find_image_rId_in_element(para._element)
        image_refs = []

        # Also check runs for images
        for run in para.runs:
            rId = find_image_rId_in_element(run._element)
            if rId and rId in image_map:
                image_refs.append(image_map[rId])

        if has_image:
            rId = has_image
            if rId in image_map:
                image_refs.append(image_map[rId])

        # Detect heading
        if para.style.name.startswith("Heading") or para.style.name.startswith("heading"):
            level_str = para.style.name.replace("Heading ", "").replace("heading ", "")
            try:
                level = int(level_str)
            except ValueError:
                level = 2
            if text:
                lines.append(f"{'#' * min(level, 4)} {text}")
            if image_refs:
                for ref in image_refs:
                    lines.append(f"![流程图](images/{ref})")
        elif text and image_refs:
            for ref in image_refs:
                lines.append(f"![流程图](images/{ref})")
            lines.append(text)
        elif text:
            lines.append(text)

        para_index += 1

    return lines


def extract_tables(doc: Document, image_map: dict) -> list[str]:
    """Extract all tables to markdown, marking images in cells."""
    lines = []
    for t_idx, table in enumerate(doc.tables):
        lines.append(f"### 表格 {t_idx + 1}")
        lines.append("")

        # Check each cell for images
        cell_images = {}
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                rId = find_image_rId_in_element(cell._element)
                if rId and rId in image_map:
                    cell_images[(r_idx, c_idx)] = image_map[rId]

        # Build table rows
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                cells.append(cell_text)
            rows.append(cells)

        if rows:
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("|" + "|".join(["---" for _ in rows[0]]) + "|")
            for r_idx, row in enumerate(rows[1:]):
                lines.append("| " + " | ".join(row) + " |")

        # Append any images found in this table
        for (img_r_idx, img_c_idx), img_name in cell_images.items():
            lines.append(f"![流程图](images/{img_name})")
            lines.append(f"*(表格{t_idx+1}中的流程图)*")

        lines.append("")

    return lines


def extract_docx(docx_path: str, output_dir: str) -> dict:
    doc = Document(docx_path)
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    img_dir = out / "images"
    img_dir.mkdir(exist_ok=True)

    # Extract ALL images from the zip
    image_map = extract_all_images_from_docx(docx_path, img_dir)

    md_lines = []
    md_lines.append(f"# 需求文档内容提取")
    md_lines.append(f"**文档路径：** {docx_path}")
    md_lines.append(f"**图片数量：** {len(image_map)}")
    if image_map:
        md_lines.append(f"**图片列表：** {', '.join(image_map.values())}")
    md_lines.append("")

    # Extract paragraphs
    md_lines.append("## 段落内容")
    md_lines.append("")
    md_lines.extend(extract_paragraphs(doc, image_map))
    md_lines.append("")

    # Extract tables
    md_lines.append("## 表格内容")
    md_lines.append("")
    md_lines.extend(extract_tables(doc, image_map))
    md_lines.append("")

    # Also check headers/footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.paragraphs:
                for para in header.paragraphs:
                    rId = find_image_rId_in_element(para._element)
                    if rId and rId in image_map:
                        md_lines.append(f"![图片 - 页眉](images/{image_map[rId]})")

        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.paragraphs:
                for para in footer.paragraphs:
                    rId = find_image_rId_in_element(footer._element)
                    if rId and rId in image_map:
                        md_lines.append(f"![图片 - 页脚](images/{image_map[rId]})")

    # Write content.md
    content_path = out / "content.md"
    content_path.write_text("\n".join(md_lines), encoding="utf-8")

    # Write metadata
    meta = {
        "docx_path": docx_path,
        "image_count": len(image_map),
        "images": list(image_map.values()),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }
    meta_path = out / "metadata.json"
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"提取完成")
    print(f"   文本内容: {content_path}")
    print(f"   图片 ({len(image_map)}张): {img_dir}")
    if image_map:
        for name in sorted(os.listdir(img_dir)):
            fpath = img_dir / name
            print(f"     - {name} ({fpath.stat().st_size} bytes)")
    print(f"   元数据: {meta_path}")

    return {
        "content_path": str(content_path),
        "images_dir": str(img_dir),
        "metadata_path": str(meta_path),
        "image_count": len(image_map),
        "images": sorted(os.listdir(img_dir)),
    }


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python extract_docx.py <需求文档.docx> <输出目录>")
        sys.exit(1)

    result = extract_docx(sys.argv[1], sys.argv[2])
    print(json.dumps(result, ensure_ascii=False))
